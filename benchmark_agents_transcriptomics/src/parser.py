"""
parser.py — read files of various types and return plain text.

Supported formats
-----------------
Plain text  : .txt .md .log .tex .xml .html .htm .py .r .rmd .sh .bash
              .sty .bbl .blg .aux (LaTeX helpers — kept as-is)
Delimited   : .csv .tsv .scv (comma-separated variant)
Excel       : .xlsx .xls .xlsm  (all sheets, up to EXCEL_MAX_ROWS_PER_SHEET)
JSON        : .json .jsonl (JSONL parsed line-by-line; agent-memory format supported)
Notebook    : .ipynb (markdown + code + cell outputs extracted)
PDF         : .pdf  (requires PyMuPDF / fitz)
Word        : .docx .doc (requires python-docx)
BibTeX      : .bib  (structured: entry type + key fields extracted)
LaTeX TOC   : .toc  (contentsline commands decoded to readable outline)
LaTeX .out  : .out  (hyperref bookmark files; UTF-16 section names decoded)
TOML        : .toml (key=value tree rendered as readable text)
Images      : .png .jpg .jpeg .gif .svg .tiff .bmp  → placeholder stub
"""

import csv
import io
import json
import logging
import re
import struct
import xml.etree.ElementTree as _ET
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Tuneable constants
# ---------------------------------------------------------------------------

# Excel: max rows / columns per sheet
EXCEL_MAX_ROWS_PER_SHEET: int = 10_000
EXCEL_MAX_COLS_PER_ROW:   int = 500

# CSV/TSV: max rows to read (prevents embedding-pipeline saturation)
CSV_MAX_ROWS: int = 10_000

# Generic line reader: minimum ratio of printable chars to keep a line,
# and minimum line length
GENERIC_MIN_PRINTABLE_RATIO: float = 0.70
GENERIC_MIN_LINE_LEN:        int   = 4

# SVG namespace
_SVG_NS = "http://www.w3.org/2000/svg"

# ---------------------------------------------------------------------------
# Optional-dependency guards
# ---------------------------------------------------------------------------

try:
    import openpyxl
    _HAS_OPENPYXL = True
except ImportError:
    _HAS_OPENPYXL = False

try:
    import pandas as _pd
    _HAS_PANDAS = True
except ImportError:
    _HAS_PANDAS = False

try:
    import fitz  # PyMuPDF
    _HAS_FITZ = True
except ImportError:
    _HAS_FITZ = False

try:
    import docx
    _HAS_DOCX = True
except ImportError:
    _HAS_DOCX = False

try:
    import tomllib          # Python 3.11+ stdlib
    _HAS_TOMLLIB = True
except ImportError:
    try:
        import tomli as tomllib  # third-party fallback
        _HAS_TOMLLIB = True
    except ImportError:
        _HAS_TOMLLIB = False


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def parse_file(filepath: Path) -> str:
    """
    Return the text content of *filepath*.
    Returns an empty string on failure.
    """
    from .platform_utils import is_ignored_file

    fp = Path(filepath)
    if is_ignored_file(fp):
        return ""

    ext = fp.suffix.lower()
    # Handle compound extensions like "results.xlsx.md" → treat as .md
    name_lower = fp.name.lower()

    try:
        # --- plain-text families ---
        if ext in {".txt", ".md", ".log", ".tex", ".xml", ".html", ".htm"}:
            return _read_plain(fp)
        elif ext in {".py", ".r", ".rmd", ".sh", ".bash"}:
            return _read_plain(fp)
        # LaTeX helper files — useful as plain text
        elif ext in {".sty", ".bbl", ".blg"}:
            return _read_plain(fp)
        elif ext == ".aux":
            return _read_aux(fp)

        # --- delimited tables ---
        elif ext in {".csv", ".tsv", ".scv"}:
            return _read_delimited(fp, ext)

        # --- Excel workbooks ---
        elif ext in {".xlsx", ".xls", ".xlsm"}:
            return _read_excel(fp)

        # --- structured text ---
        elif ext == ".json":
            return _read_json(fp)
        elif ext == ".jsonl":
            return _read_jsonl(fp)
        elif ext == ".toml":
            return _read_toml(fp)

        # --- notebooks ---
        elif ext == ".ipynb":
            return _read_notebook(fp)

        # --- binary documents ---
        elif ext == ".pdf":
            return _read_pdf(fp)
        elif ext in {".docx", ".doc"}:
            return _read_docx(fp)

        # --- BibTeX / LaTeX bibliography ---
        elif ext == ".bib":
            return _read_bib(fp)

        # --- LaTeX outline/TOC files ---
        elif ext == ".toc":
            return _read_toc(fp)

        # --- LaTeX hyperref bookmark file ---
        elif ext == ".out":
            return _read_latex_out(fp)

        # --- SVG: XML vector format — extract embedded text ---
        elif ext == ".svg":
            return _read_svg(fp)

        # --- raster images: no text without OCR — return informative stub ---
        elif ext in {".png", ".jpg", ".jpeg", ".gif", ".tiff", ".bmp"}:
            return f"[IMAGE FILE: {fp.name}]"

        # --- binary numeric arrays — stub to avoid flooding embeddings ---
        elif ext in {".parquet", ".npy", ".feather", ".rds"}:
            size_mb = fp.stat().st_size / 1_048_576
            return f"[BINARY DATA FILE: {fp.name}, {size_mb:.1f} MB]"

        # --- macOS metadata — skip silently ---
        elif name_lower.endswith(".ds_store"):
            return ""

        # --- catch-all: try plain text first, then line-by-line fallback ---
        else:
            try:
                text = _read_plain(fp)
                if text.strip():
                    return text
            except Exception:
                pass
            return _read_generic_lines(fp)

    except Exception as exc:
        logger.warning("Failed to parse %s: %s", fp, exc)
        return ""


# ---------------------------------------------------------------------------
# Plain text
# ---------------------------------------------------------------------------

def _read_plain(fp: Path) -> str:
    for enc in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
        try:
            return fp.read_text(encoding=enc)
        except (UnicodeDecodeError, UnicodeError):
            continue
    return fp.read_bytes().decode("utf-8", errors="replace")


# ---------------------------------------------------------------------------
# Delimited (.csv / .tsv / .scv) — line-by-line with row cap + auto-sniff
# ---------------------------------------------------------------------------

def _read_delimited(fp: Path, ext: str) -> str:
    """
    Parse a delimited file line by line without loading the whole file into
    memory first.  Separator is inferred from the extension; for ambiguous
    cases the csv.Sniffer is used on the first 4 KB.

    Rows are capped at CSV_MAX_ROWS to prevent huge gene-count tables from
    flooding the embedding pipeline with repetitive numeric tokens.
    """
    default_sep = "\t" if ext == ".tsv" else ","

    for enc in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
        try:
            lines_out: list[str] = []
            n_rows = 0
            header_raw = ""
            sep = default_sep

            with open(fp, newline="", encoding=enc, errors="replace") as fh:
                # Try to sniff separator from first 4 KB
                sample = fh.read(4096)
                fh.seek(0)
                try:
                    dialect = csv.Sniffer().sniff(sample, delimiters=",\t;|")
                    sep = dialect.delimiter
                except csv.Error:
                    sep = default_sep

                reader = csv.reader(fh, delimiter=sep)
                for row in reader:
                    if n_rows >= CSV_MAX_ROWS:
                        lines_out.append(
                            f"[... truncated after {CSV_MAX_ROWS} rows ...]"
                        )
                        break
                    lines_out.append("\t".join(row))
                    n_rows += 1

            logger.debug("CSV %s: %d rows, sep=%r", fp.name, n_rows, sep)
            return "\n".join(lines_out)
        except (UnicodeDecodeError, UnicodeError):
            continue
        except Exception as exc:
            logger.warning("Delimited parse error %s: %s", fp.name, exc)
            return _read_plain(fp)

    return _read_plain(fp)


# ---------------------------------------------------------------------------
# SVG (.svg) — XML vector graphics; extract all embedded text
# ---------------------------------------------------------------------------

# Regex to identify Matplotlib/Inkscape font glyph symbol IDs like:
#   DejaVuSans-41  (U+0041 = 'A')
#   DejaVuSans-Bold-70 (U+0070 = 'p')
_SVG_GLYPH_ID_RE = re.compile(r"[A-Za-z][A-Za-z\-]+?-([0-9A-Fa-f]{1,5})$")

# Structural element ID words worth reporting (Matplotlib naming conventions)
_SVG_STRUCTURAL_KEYWORDS = {
    "axes", "figure", "xtick", "ytick", "xlabel", "ylabel",
    "title", "legend", "line", "text", "patch", "bar", "scatter",
    "image", "colorbar", "annotation",
}


def _decode_glyph_ids(root: _ET.Element) -> str:
    """
    For vectorized SVGs (text-as-paths): decode font glyph symbol IDs in
    <defs> to reconstruct the set of characters used in the figure.

    Example: DejaVuSans-50 → 'P', DejaVuSans-43 → 'C', etc.
    Returns a string like "Characters used: P C A 1 2 ..." or "" if none.
    """
    chars: list[str] = []
    for elem in root.iter():
        eid = elem.get("id") or ""
        m = _SVG_GLYPH_ID_RE.match(eid)
        if m:
            try:
                codepoint = int(m.group(1), 16)
                ch = chr(codepoint)
                if ch.isprintable() and not ch.isspace():
                    chars.append(ch)
            except (ValueError, OverflowError):
                pass
    if not chars:
        return ""
    # Remove duplicates, sort meaningful chars first
    seen: set[str] = set()
    unique: list[str] = []
    for c in chars:
        if c not in seen:
            seen.add(c)
            unique.append(c)
    return "Characters used in figure: " + " ".join(unique)


def _read_svg(fp: Path) -> str:
    """
    Extract human-readable text from an SVG file.

    Two SVG modes are handled:

    A. Text-as-XML (Inkscape, hand-written SVGs):
       Extracts <text>, <tspan>, <textPath>, <title>, <desc>, and
       aria-label / title attributes.

    B. Text-as-paths (Matplotlib with svg.fonttype=3, default):
       No <text> elements exist.  Instead we:
         1. Decode font glyph symbol IDs from <defs> to recover the
            character set (e.g. DejaVuSans-50 → 'P').
         2. Report structural element IDs (axes_1, xtick_1, xlabel, etc.)
            which encode the figure's structural composition.
         3. Extract the figure filename as a semantic hint.
         4. Extract <metadata> (creator, date).

    Falls back to _read_generic_lines if XML parsing fails entirely.
    """
    try:
        raw = _read_plain(fp)
        raw_clean = re.sub(r"<!DOCTYPE[^>]*>", "", raw, flags=re.DOTALL)
        root = _ET.fromstring(raw_clean)

        def _local(tag: str) -> str:
            return tag.split("}")[-1] if "}" in tag else tag

        def _all_text(el: _ET.Element) -> str:
            return " ".join(t.strip() for t in el.itertext() if t and t.strip())

        parts: list[str] = [f"[SVG FILE: {fp.name}]"]

        # --- 1. Document-level <title> and <desc> ---
        seen: set[str] = set()
        for tag in ("title", "desc"):
            el = root.find(f"{{{_SVG_NS}}}{tag}") or root.find(tag)
            if el is not None:
                t = _all_text(el)
                if t and t not in seen:
                    seen.add(t)
                    parts.append(f"{tag.upper()}: {t}")

        # --- 2. All <text>/<tspan>/<textPath> elements (Mode A) ---
        text_parts: list[str] = []
        label_parts: list[str] = []

        for elem in root.iter():
            local = _local(elem.tag)
            if local in ("text", "tspan", "textPath"):
                t = _all_text(elem)
                if t and t not in seen:
                    seen.add(t)
                    text_parts.append(t)
            elif local == "title":
                t = _all_text(elem)
                if t and t not in seen:
                    seen.add(t)
                    label_parts.append(t)
            for attr in ("aria-label",):
                val = elem.get(attr)
                if val and val.strip() and val.strip() not in seen:
                    seen.add(val.strip())
                    label_parts.append(val.strip())

        if text_parts:
            parts.append(f"\nText elements ({len(text_parts)}):")
            parts.extend(text_parts)
        if label_parts:
            parts.append(f"\nLabels / tooltips ({len(label_parts)}):")
            parts.extend(label_parts)

        # --- 3. Mode B: vectorized text — glyph decoding + structural IDs ---
        if not text_parts:
            glyph_str = _decode_glyph_ids(root)
            if glyph_str:
                parts.append(f"\n{glyph_str}")

            # Structural element IDs
            struct_ids: list[str] = []
            seen_ids: set[str] = set()
            for elem in root.iter():
                eid = elem.get("id") or ""
                if not eid or _SVG_GLYPH_ID_RE.match(eid):
                    continue
                # Keep IDs that contain a known structural keyword
                eid_lower = eid.lower()
                if any(k in eid_lower for k in _SVG_STRUCTURAL_KEYWORDS):
                    if eid not in seen_ids:
                        seen_ids.add(eid)
                        struct_ids.append(eid)
            if struct_ids:
                parts.append(f"\nFigure structure elements ({len(struct_ids)}):")
                parts.append(", ".join(struct_ids))

            # Filename as semantic hint (e.g. "fig2_volcano" → volcano plot)
            stem = fp.stem
            hint = re.sub(r"[-_]+", " ", stem).strip()
            if hint:
                parts.append(f"\nFilename hint: {hint}")

        # --- 4. Metadata ---
        for meta in root.iter(f"{{{_SVG_NS}}}metadata"):
            t = _all_text(meta)
            if t:
                parts.append(f"\nMetadata: {t}")
                break

        result = "\n".join(parts)
        logger.info(
            "SVG %s: %d text elements, %d labels",
            fp.name, len(text_parts), len(label_parts),
        )
        return result

    except Exception as exc:
        logger.warning("SVG XML parse error %s: %s — trying generic reader", fp.name, exc)
        return _read_generic_lines(fp)


# ---------------------------------------------------------------------------
# Generic line-by-line fallback for any file (binary or text)
# ---------------------------------------------------------------------------

def _read_generic_lines(
    fp: Path,
    min_printable_ratio: float = GENERIC_MIN_PRINTABLE_RATIO,
    min_len: int = GENERIC_MIN_LINE_LEN,
) -> str:
    """
    Last-resort text extractor for any file.

    Reads raw bytes, splits on newline, and keeps every line where at least
    `min_printable_ratio` of non-whitespace characters are printable ASCII/Unicode.
    This recovers readable strings from:
      - Partially-binary files with embedded text
      - Files with unexpected encodings
      - Format stubs that fell through all specific parsers

    Returns a stub message if nothing readable is found.
    """
    try:
        raw = fp.read_bytes()
        out: list[str] = []
        for raw_line in raw.split(b"\n"):
            try:
                line = raw_line.decode("utf-8", errors="replace").rstrip()
            except Exception:
                continue
            if len(line) < min_len:
                continue
            non_ws = line.replace(" ", "").replace("\t", "")
            if not non_ws:
                continue
            n_print = sum(1 for c in non_ws if c.isprintable())
            if n_print / len(non_ws) >= min_printable_ratio:
                out.append(line)

        if out:
            logger.debug("Generic reader %s: %d readable lines", fp.name, len(out))
            return "\n".join(out)
        return f"[BINARY FILE: {fp.name} — no readable text found]"
    except Exception as exc:
        logger.warning("Generic line reader failed for %s: %s", fp.name, exc)
        return ""


# ---------------------------------------------------------------------------
# Excel (.xlsx / .xlsm / .xls) — all sheets
# ---------------------------------------------------------------------------

def _read_excel(fp: Path) -> str:
    """
    Parse an Excel workbook and return all sheets as labelled plain text.

    Strategy (in order of preference):
    1. openpyxl   — .xlsx and .xlsm
    2. pandas     — .xls and fallback for .xlsx
    3. Stub       — if neither library is available

    Safety limits: EXCEL_MAX_ROWS_PER_SHEET rows, EXCEL_MAX_COLS_PER_ROW columns.
    Fully-empty rows are skipped.
    """
    ext = fp.suffix.lower()
    parts: list[str] = []

    # --- attempt 1: openpyxl (.xlsx / .xlsm) ---
    if _HAS_OPENPYXL and ext in {".xlsx", ".xlsm"}:
        try:
            wb = openpyxl.load_workbook(str(fp), read_only=True, data_only=True)
            sheet_names = wb.sheetnames
            parts.append(
                f"[WORKBOOK: {fp.name} | sheets={len(sheet_names)} | "
                f"names={', '.join(sheet_names)}]"
            )
            for sheet_name in sheet_names:
                ws = wb[sheet_name]
                parts.append(f"\n=== Sheet: {sheet_name} ===")
                n_rows_written = 0
                for row in ws.iter_rows(values_only=True):
                    if all(c is None for c in row):
                        continue
                    if n_rows_written >= EXCEL_MAX_ROWS_PER_SHEET:
                        parts.append(
                            f"[... truncated after {EXCEL_MAX_ROWS_PER_SHEET} rows ...]"
                        )
                        break
                    cells = [str(c) if c is not None else ""
                             for c in row[:EXCEL_MAX_COLS_PER_ROW]]
                    parts.append("\t".join(cells))
                    n_rows_written += 1
                logger.debug("Sheet '%s' in %s: %d rows", sheet_name, fp.name, n_rows_written)
            wb.close()
            logger.info("Parsed %d sheet(s) from %s", len(sheet_names), fp.name)
            return "\n".join(parts)
        except Exception as exc:
            logger.warning("openpyxl failed on %s (%s); trying pandas", fp.name, exc)
            parts = []

    # --- attempt 2: pandas fallback ---
    if _HAS_PANDAS:
        try:
            xl = _pd.ExcelFile(str(fp))
            sheet_names = xl.sheet_names
            parts.append(
                f"[WORKBOOK: {fp.name} | sheets={len(sheet_names)} | "
                f"names={', '.join(str(s) for s in sheet_names)}]"
            )
            for sheet_name in sheet_names:
                df = xl.parse(
                    sheet_name, header=0,
                    nrows=EXCEL_MAX_ROWS_PER_SHEET, dtype=str,
                ).fillna("")
                if df.shape[1] > EXCEL_MAX_COLS_PER_ROW:
                    df = df.iloc[:, :EXCEL_MAX_COLS_PER_ROW]
                parts.append(f"\n=== Sheet: {sheet_name} ===")
                parts.append("\t".join(str(c) for c in df.columns))
                for _, row in df.iterrows():
                    row_vals = list(row.astype(str))
                    if all(v.strip() in ("", "nan") for v in row_vals):
                        continue
                    parts.append("\t".join(row_vals))
            xl.close()
            logger.info("Parsed %d sheet(s) from %s (pandas)", len(sheet_names), fp.name)
            return "\n".join(parts)
        except Exception as exc:
            logger.warning("pandas Excel parse error %s: %s", fp.name, exc)

    # --- attempt 3: stub ---
    missing = []
    if not _HAS_OPENPYXL:
        missing.append("openpyxl")
    if not _HAS_PANDAS:
        missing.append("pandas")
    hint = f"install {' or '.join(missing)} to parse" if missing else "file may be corrupt"
    logger.warning("Could not parse Excel file %s — %s", fp.name, hint)
    return f"[EXCEL FILE: {fp.name} — {hint}]"


# ---------------------------------------------------------------------------
# JSON (.json)
# ---------------------------------------------------------------------------

def _read_json(fp: Path) -> str:
    try:
        data = json.loads(_read_plain(fp))
        return json.dumps(data, indent=2, ensure_ascii=False)
    except Exception:
        return _read_plain(fp)


# ---------------------------------------------------------------------------
# JSONL (.jsonl) — one JSON object per line
# ---------------------------------------------------------------------------

# Fields extracted from agent-memory / transcript JSONL files
_JSONL_CONTENT_KEYS = ("content", "text", "message", "value", "output", "result")
_JSONL_ROLE_KEYS    = ("role", "type", "author")


def _extract_jsonl_text(obj: dict) -> str:
    """Pull the most informative text out of one parsed JSONL object."""
    parts = []

    # Support Finch/Heron agent-memory format:
    # {"ts":..., "step":..., "message": {"role": "user"|"assistant", "content": "..."}}
    msg = obj.get("message")
    if isinstance(msg, dict):
        role    = msg.get("role", "")
        content = msg.get("content", "")
        if isinstance(content, str) and content.strip():
            prefix = f"[{role}] " if role else ""
            return prefix + content.strip()
        elif isinstance(content, list):
            # content can be a list of {"type":"text","text":"..."} blocks
            texts = []
            for block in content:
                if isinstance(block, dict):
                    texts.append(block.get("text") or block.get("content") or "")
                elif isinstance(block, str):
                    texts.append(block)
            combined = " ".join(t for t in texts if t.strip())
            if combined:
                prefix = f"[{role}] " if role else ""
                return prefix + combined.strip()

    # Generic fallback: look for common content/role keys
    role = ""
    for k in _JSONL_ROLE_KEYS:
        if k in obj and isinstance(obj[k], str):
            role = obj[k]
            break
    text = ""
    for k in _JSONL_CONTENT_KEYS:
        if k in obj:
            v = obj[k]
            if isinstance(v, str) and v.strip():
                text = v.strip()
                break
            elif isinstance(v, (dict, list)):
                text = json.dumps(v, ensure_ascii=False)
                break

    if text:
        return f"[{role}] {text}" if role else text

    # Last resort: dump the whole object compactly
    return json.dumps(obj, ensure_ascii=False, separators=(",", ":"))


def _read_jsonl(fp: Path) -> str:
    """
    Parse a JSONL file (one JSON object per line).

    For Finch / Heron agent-memory files the format is:
      {"ts": ..., "step": N, "message": {"role": "...", "content": "..."}}

    For each line the most meaningful text is extracted (role + content).
    Non-JSON lines and error lines are kept as-is.
    Lines whose content is purely metadata with no human-readable text are
    skipped to reduce embedding noise.
    """
    raw = _read_plain(fp)
    lines = raw.splitlines()
    parts: list[str] = [f"[JSONL FILE: {fp.name} | lines={len(lines)}]"]
    n_parsed = 0
    n_skipped = 0

    for i, line in enumerate(lines):
        line = line.strip()
        if not line:
            continue
        try:
            obj = json.loads(line)
            text = _extract_jsonl_text(obj) if isinstance(obj, dict) else str(obj)
            if text.strip():
                parts.append(text)
                n_parsed += 1
            else:
                n_skipped += 1
        except json.JSONDecodeError:
            parts.append(line)  # keep raw non-JSON lines (e.g. log messages)

    logger.info(
        "JSONL %s: %d lines parsed, %d skipped (empty/metadata)",
        fp.name, n_parsed, n_skipped,
    )
    return "\n\n".join(parts)


# ---------------------------------------------------------------------------
# TOML (.toml)
# ---------------------------------------------------------------------------

def _toml_flatten(obj: object, prefix: str = "") -> list[str]:
    """Recursively render a TOML value tree as 'key = value' lines."""
    lines: list[str] = []
    if isinstance(obj, dict):
        for k, v in obj.items():
            full_key = f"{prefix}.{k}" if prefix else k
            if isinstance(v, (dict, list)):
                lines.extend(_toml_flatten(v, full_key))
            else:
                lines.append(f"{full_key} = {v!r}")
    elif isinstance(obj, list):
        for idx, item in enumerate(obj):
            lines.extend(_toml_flatten(item, f"{prefix}[{idx}]"))
    else:
        lines.append(f"{prefix} = {obj!r}")
    return lines


def _read_toml(fp: Path) -> str:
    """
    Parse a TOML file and render it as flat key=value text.
    Falls back to plain-text reading if tomllib / tomli is unavailable.
    """
    if not _HAS_TOMLLIB:
        logger.debug("tomllib/tomli not available; reading %s as plain text", fp.name)
        return _read_plain(fp)
    try:
        with open(fp, "rb") as fh:
            data = tomllib.load(fh)   # type: ignore[attr-defined]
        lines = [f"[TOML FILE: {fp.name}]"] + _toml_flatten(data)
        return "\n".join(lines)
    except Exception as exc:
        logger.warning("TOML parse error %s: %s — falling back to plain text", fp.name, exc)
        return _read_plain(fp)


# ---------------------------------------------------------------------------
# Jupyter Notebook (.ipynb)
# ---------------------------------------------------------------------------

def _read_notebook(fp: Path) -> str:
    try:
        nb = json.loads(_read_plain(fp))
        parts = []
        for cell in nb.get("cells", []):
            ct  = cell.get("cell_type", "")
            src = cell.get("source", [])
            text = "".join(src) if isinstance(src, list) else str(src)
            if ct == "markdown":
                parts.append(f"[MARKDOWN]\n{text}")
            elif ct == "code":
                parts.append(f"[CODE]\n{text}")
                for out in cell.get("outputs", []):
                    if "text" in out:
                        parts.append("[OUTPUT]\n" + "".join(out["text"]))
            else:
                parts.append(text)
        return "\n\n".join(parts)
    except Exception:
        return _read_plain(fp)


# ---------------------------------------------------------------------------
# PDF (.pdf)
# ---------------------------------------------------------------------------

def _read_pdf(fp: Path) -> str:
    if not _HAS_FITZ:
        logger.debug("PyMuPDF not available; skipping %s", fp.name)
        return f"[PDF FILE: {fp.name} — install PyMuPDF to parse]"
    try:
        doc = fitz.open(str(fp))
        pages = [page.get_text() for page in doc]
        doc.close()
        return "\n".join(pages)
    except Exception as exc:
        logger.warning("PDF parse error %s: %s", fp, exc)
        return f"[PDF FILE: {fp.name}]"


# ---------------------------------------------------------------------------
# Word (.docx / .doc)
# ---------------------------------------------------------------------------

def _read_docx(fp: Path) -> str:
    if not _HAS_DOCX:
        logger.debug("python-docx not available; skipping %s", fp.name)
        return f"[DOCX FILE: {fp.name} — install python-docx to parse]"
    try:
        doc = docx.Document(str(fp))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = "\t".join(c.text.strip() for c in row.cells if c.text.strip())
                if row_text:
                    parts.append(row_text)
        return "\n".join(parts)
    except Exception as exc:
        logger.warning("DOCX parse error %s: %s", fp, exc)
        return f"[DOCX FILE: {fp.name}]"


# ---------------------------------------------------------------------------
# BibTeX (.bib)
# ---------------------------------------------------------------------------

# Matches a full BibTeX entry: @type{key, ...}
_BIB_ENTRY_RE  = re.compile(
    r"@(\w+)\s*\{\s*([^,\s]+)\s*,\s*(.*?)\n\}",
    re.DOTALL | re.IGNORECASE,
)
# Matches individual fields inside an entry: fieldname = {value} or "value" or number
_BIB_FIELD_RE  = re.compile(
    r"(\w+)\s*=\s*(?:\{([^{}]*(?:\{[^{}]*\}[^{}]*)*)\}|\"([^\"]*)\"|(\d+))",
    re.DOTALL,
)
# Fields to include in the output (ordered by relevance for embeddings)
_BIB_PRIORITY_FIELDS = [
    "title", "author", "year", "journal", "booktitle",
    "abstract", "keywords", "doi",
]


def _read_bib(fp: Path) -> str:
    """
    Parse a BibTeX (.bib) file and extract structured entry information.
    Returns a plain-text block per entry with labelled fields.
    Falls back to raw text if parsing fails.
    """
    raw = _read_plain(fp)
    entries = _BIB_ENTRY_RE.findall(raw)
    if not entries:
        return raw  # nothing matched — return as-is

    parts = [f"[BIBTEX FILE: {fp.name} | entries={len(entries)}]"]
    for entry_type, key, body in entries:
        lines = [f"\n@{entry_type.upper()} {{{key}}}"]
        fields: dict[str, str] = {}
        for fname, fval_brace, fval_quote, fval_num in _BIB_FIELD_RE.findall(body):
            value = fval_brace or fval_quote or fval_num
            fields[fname.lower()] = value.strip().replace("\n", " ")
        # Output priority fields first, then the rest
        seen = set()
        for f in _BIB_PRIORITY_FIELDS:
            if f in fields:
                lines.append(f"  {f}: {fields[f]}")
                seen.add(f)
        for f, v in fields.items():
            if f not in seen:
                lines.append(f"  {f}: {v}")
        parts.append("\n".join(lines))

    return "\n".join(parts)


# ---------------------------------------------------------------------------
# LaTeX auxiliary file (.aux)
# ---------------------------------------------------------------------------

# Pull citation keys and cross-reference labels — the parts useful for analysis
_AUX_CITATION_RE = re.compile(r"\\citation\{([^}]+)\}")
_AUX_NEWLABEL_RE = re.compile(r"\\newlabel\{([^}]+)\}\{\{([^}]*)\}\{([^}]*)\}")


def _read_aux(fp: Path) -> str:
    """
    Extract citations and section labels from a LaTeX .aux file.
    Falls back to plain-text reading if no structured content is found.
    """
    raw = _read_plain(fp)
    citations = _AUX_CITATION_RE.findall(raw)
    labels    = _AUX_NEWLABEL_RE.findall(raw)

    if not citations and not labels:
        return raw

    parts = [f"[LATEX AUX: {fp.name}]"]
    if citations:
        unique_cites = sorted(set(c.strip() for cites in citations for c in cites.split(",")))
        parts.append(f"Citations ({len(unique_cites)}): {', '.join(unique_cites)}")
    if labels:
        parts.append(f"\nCross-reference labels ({len(labels)}):")
        for name, page, anchor in labels[:100]:   # cap at 100 for very large docs
            parts.append(f"  {name} → page {page}")
    return "\n".join(parts)


# ---------------------------------------------------------------------------
# LaTeX table-of-contents (.toc)
# ---------------------------------------------------------------------------

# \contentsline{section}{\numberline {1}Title}{page}{anchor}
# \contentsline{section}{Title}{page}{anchor}
_TOC_CONTENTSLINE_RE = re.compile(
    r"\\contentsline\s*\{(\w+)\}\s*\{(?:\\numberline\s*\{[^}]*\}\s*)?([^}]+)\}\s*\{(\d+)\}",
    re.DOTALL,
)
# Depth ordering for indentation
_TOC_DEPTH = {
    "part": 0, "chapter": 1, "section": 1, "subsection": 2,
    "subsubsection": 3, "paragraph": 4,
}


def _read_toc(fp: Path) -> str:
    """
    Parse a LaTeX .toc file into a human-readable section outline.
    LaTeX commands inside section titles are stripped.
    """
    raw = _read_plain(fp)
    entries = _TOC_CONTENTSLINE_RE.findall(raw)
    if not entries:
        return raw

    parts = [f"[LATEX TOC: {fp.name} | entries={len(entries)}]"]
    for level, title, page in entries:
        # Strip remaining LaTeX commands from title
        title_clean = re.sub(r"\\[a-zA-Z]+\s*\{([^}]*)\}", r"\1", title)
        title_clean = re.sub(r"\\[a-zA-Z]+\s*", "", title_clean).strip()
        depth  = _TOC_DEPTH.get(level.lower(), 2)
        indent = "  " * depth
        parts.append(f"{indent}{level}: {title_clean} (p.{page})")
    return "\n".join(parts)


# ---------------------------------------------------------------------------
# LaTeX hyperref bookmarks (.out)
# ---------------------------------------------------------------------------

# Matches: \BOOKMARK [depth][-]{anchor}{\376\377\000X\000Y...}{parent}% num
_OUT_BOOKMARK_RE = re.compile(
    r"\\BOOKMARK\s*\[(-?\d+)\]\[-\]\{[^}]*\}\{(.*?)\}\{[^}]*\}",
    re.DOTALL,
)
# Matches the LaTeX octal-escape for a UTF-16 BOM followed by encoded chars
_UTF16_CHAR_RE = re.compile(r"\\000(.)")   # \000X → char X (BMP)


def _decode_bookmark_title(raw_title: str) -> str:
    """
    Decode LaTeX UTF-16BE bookmark encoding to a plain string.

    The hyperref package stores non-ASCII titles as UTF-16BE bytes written as
    LaTeX octal escapes: \\376\\377 (BOM) followed by \\000X for each BMP char.
    """
    # Remove the UTF-16 BOM escape if present
    cleaned = raw_title.replace(r"\376\377", "").replace("\\376\\377", "")
    # Decode \000X sequences → char X
    decoded = _UTF16_CHAR_RE.sub(lambda m: m.group(1), cleaned)
    # Strip any remaining backslash-commands
    decoded = re.sub(r"\\[a-zA-Z]+\s*", "", decoded)
    # Replace \040 (LaTeX octal for space)
    decoded = decoded.replace(r"\040", " ").replace("\\040", " ")
    return decoded.strip()


def _read_latex_out(fp: Path) -> str:
    """
    Parse a LaTeX hyperref .out (PDF bookmarks) file.

    The file contains \\BOOKMARK entries with section titles encoded as
    UTF-16BE octal sequences.  We decode them to readable text and render
    the document outline.  Falls back to raw text for non-hyperref .out files.
    """
    raw = _read_plain(fp)
    entries = _OUT_BOOKMARK_RE.findall(raw)
    if not entries:
        # Not a hyperref bookmarks file — return as plain text
        return raw

    parts = [f"[LATEX BOOKMARKS: {fp.name} | entries={len(entries)}]",
             "Document outline:"]
    for depth_str, raw_title in entries:
        depth = int(depth_str)
        title = _decode_bookmark_title(raw_title)
        if title:
            indent = "  " * max(depth, 0)
            parts.append(f"{indent}{title}")
    return "\n".join(parts)
